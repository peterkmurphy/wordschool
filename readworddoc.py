# Used for writing student reports.

from docx import Document
import yaml
import os
import glob
import errno

# From https://stackoverflow.com/questions/273192/how-can-i-create-a-directory-if-it-does-not-exist
def make_sure_path_exists(path):
    try:
        os.makedirs(path)
    except OSError as exception:
        if exception.errno != errno.EEXIST:
            raise


class ReportWriter:
    ''' This is for writing student reports using Word. '''

    def __init__(self, filename):
        ''' We initialise this using a existing file as a template. '''
        self.filename = filename
        self.document = Document(filename)
        self.heading = self.document.tables[0]
        self.marks = self.document.tables[1]

    def assigntexttotable(self, table, row, col, parnum, text):
        ''' Writes text to a cell of a table in a Word doc. Arguments:
        table: the table in a Word document
        row: the row number (zero based)
        col: the col number (zero based)
        parnum: the number of the paragraph.
        text: the text to write.

        Note: if there is existing text at the destination paragraph, then the
        method overwrites the text while keeping the existing font. However,
        if there is no text in the paragraph, then the method just adds text
        to the paragraph; the font defaults to the "Normal" style, which is
        perhaps not what you want. This is probably a quirk of python-docx
        which needs to be worked around.
        '''
        runstocount = len(table.cell(row,col).paragraphs[parnum].runs)
        if (runstocount > 0):
            table.cell(row,col).paragraphs[parnum].runs[0].text = text
            for i in range(runstocount - 1):
                table.cell(row,col).paragraphs[parnum].runs[i + 1].clear()
        else:
            table.cell(row,col).paragraphs[parnum].add_run(text)

    def writeinitial(self, initid):
        ''' This writes the report using a dict (initid) with the
        following possible keys:
        id: the student id
        name: the name of the student
        course: the course (like "General English")
        cl: the class (like "Pre-Intermediate")
        sd: the starting date
        ed: the end date.
        '''

        if initid.has_key("id"):
            self.assigntexttotable(self.heading, 0, 0, 2, initid["id"])
        if initid.has_key("name"):
            self.assigntexttotable(self.heading, 0, 1, 2, initid["name"])
        if initid.has_key("course"):
            self.assigntexttotable(self.heading, 0, 2, 1, initid["course"])
        if initid.has_key("cl"):
            self.assigntexttotable(self.heading, 1, 2, 1, initid["cl"])
        if initid.has_key("sd"):
            self.assigntexttotable(self.heading, 0, 3, 2, initid["sd"])
        if initid.has_key("ed"):
            self.assigntexttotable(self.heading, 0, 4, 2, initid["ed"])

    def writedates(self, weekdates):
        ''' Writes the starting dates for each week of class to the report.
        Since there are 10 weeks, there should be ten values in weekdates.
        By convention, they should be the starting Monday for each week.

        Note: there are two "Classroom Progress" rows in the report, and so we
        don't write the dates for those row.
        '''
        for i in range(5):
            self.assigntexttotable(self.marks, i + 1, 1, 0, weekdates[i])
        for i in range(5):
            self.assigntexttotable(self.marks, i + 7, 1, 0, weekdates[5 + i])

    def writemarks(self, firstweek, lastweek, marks):
        ''' This writes marks to a table of marks, under the assumption that:
        (a) There are 10 weeks to write to, and 6 criteria to mark by;
        each row is for a week, and each column is a criterion.
        (b) There is a progress report for week 5 (with 6 criteria to mark by);
        if you write marks for week 5, you need to mark the subsequent
        progress report row.
        (c) There is a progress report for week 10 (with 6 criteria to mark by);
        if you write marks for week 10, you need to mark the subsequent
        progress report row.
        There is a firstweek and a lastweek parameter, with the condition:
        1 <= firstweek <= lastweek <= 10
        There is also a marks table, which contains the marks for each week
        (and progress reports as necessary). For each element marks[i][j],
        i represents row indices, and j represents criteria in columns.
        '''
        numweeks = lastweek - firstweek + 1
        if firstweek <= 5 and lastweek <= 5:
            numweeks += 1
        if lastweek == 10:
            numweeks += 1
        if firstweek <= 5:
            startindex = firstweek
        else:
            startindex = firstweek + 1
        for i in range(len(marks)):
            for j in range(len(marks[i])):
                self.assigntexttotable(self.marks, startindex + i, j + 2, 0,
                    marks[i][j])

    def writecomment(self, comment):
        ''' This writes a comment for a student report. '''
        self.assigntexttotable(self.marks, 13, 0, 1, "Comment: " + comment)

    def save(self, newfilename = None):
        ''' This saves the document to a new file (newfilename). If no argument
        is present, information is saved to the original source.
        '''
        if newfilename:
            self.document.save(newfilename)
        else:
            self.document.save(self.filename)

    @staticmethod
    def WriteReports(mode, reportyaml, outputdir, isgradsubdir = False,
        graddir="Graduated", template=None):
        ''' This is the main method for the ReportWriter class. Its purpose,
        as the name says, is to write reports, based on data in a YAML file,
        to Word documents in a directory. The WriteReports takes the following
        arguments:

        mode: what is the main goal to be achived in using the function? The
        modes are:
            - "N": for new: takes a template, writes the report data from a
            YAML file to it, and saves it as a file to the output directory. The
            file name of the report is always derived from the name of the
            student (which is in the report). This mode may overwrite existing
            reports, so use with caution. However, this mode is probably suited
            for a new class, where reports need to created from scratch.
            - "E": for existing: looks for student names in the report as
            above, and looks for files in the output directory matching those
            names. When such files are found, data from the YAML file is
            appended to the file, but existing data is not overwritten. This
            mode is probably suited for an existing class with existing
            class records.
            - "T": tidy the dates. Looks for all Word docs in a directory, and
            sets the term dates to new values specifed in the YAML file. This
            mode is used when there are existing reports that need to be reused
            for a new term.
        reportyaml: the YAML file that contains the report data. This should
        consist of several YAML documents. If there is data specifying the term
        dates, this should be the first document inside. (This function can
        automatically check if there is term date data.) Subsequent documents
        should all correspond to student data, with each student represented
        by a single YAML document inside the file.
        outputdir: the name of the directory where reports need to be written.
        This could be an absolute path, or one relative to the current working
        directory.
        isgradsubdir: whether reports for graduate students are written to a
        subdirectory of outputdir (or not). The default is False. (See below on
        what indicates a students is a "graduate".) This setting is ignored for
        "T" mode.
        graddir: the name of the subdirectory for reports for graduate students.
        The default is "Graduated". Like above, this setting is ignored for
        "T" mode.
        template: the file name of a .docx word which can be used as a template
        for creating templates. This is used with "N" (new) mode.

        Note: for the purpose of this program, students with non-empty comment
        information are considered "graduates", and students with empty comment
        information are not considered graduates.
        '''

        make_sure_path_exists(outputdir)
        if isgradsubdir:
            make_sure_path_exists(os.path.join(outputdir, graddir))

        reportdata = list(yaml.load_all(open(reportyaml, 'r')))

# Check for date data and student data (if either thing exists)

        if isinstance(reportdata[0][0], dict):
            thedates = None
            thestudents = reportdata
        else:
            thedates = reportdata[0]
            thestudents = reportdata[1:]

# Are we dealing with tidy mode?

        if mode == 'T':
            os.chdir(outputdir)
            for existreport in glob.glob("*.docx"):
                R = ReportWriter(existreport)
                if thedates:
                    R.writedates(thedates)
                R.save()

        else:
            for student in thestudents:
                basename = student[0]["name"]+".docx"
                hascomment = False
                if len(student) > 1 and student[1]["comment"]:
                    hascomment = True
                if hascomment and isgradsubdir:
                    studentname = os.path.join(outputdir, graddir, basename)
                else:
                    studentname = os.path.join(outputdir, basename)
                if mode == "N":
                    R = ReportWriter(template)
                else:
                    R = ReportWriter(studentname)
                if hascomment:
                    R.writecomment(student[1]["comment"])
                if len(student) > 2:
                    R.writemarks(student[2]["start"], student[2]["end"],
                        student[3:])
                R.writeinitial(student[0])
                if thedates:
                    R.writedates(thedates)
                if mode == "N":
                    R.save(studentname)
                else:
                    R.save()

# This is where the command line arguments are read and understood.

import argparse
parser = argparse.ArgumentParser(
    description='Generate student reports in Word docx form from YAML')
parser.add_argument('mode',
    help='N (creating New docs); E (updating Existing docs); T (tidy dates)')
parser.add_argument('reportyaml',
    help='A YAML file containing student data')
parser.add_argument('outputdir',
    help='The output directory for generated student reports')
parser.add_argument('--gradsub', dest='isgradsubdir', action='store_true',
    help='Graduate students reports written to subdirectory of the output dir')
parser.add_argument('--no-gradsub', dest='isgradsubdir', action='store_false',
    help='Graduate students reports written directly to output dir')
parser.set_defaults(isgradsubdir=True)
parser.add_argument('--graddir', default="Graduated",
    help='Subdirectory for graduate students (default: "Graduated")')
parser.add_argument('--template', default="None",
    help='Template file in docx format used for generating reports')
args = parser.parse_args()
ReportWriter.WriteReports(args.mode, args.reportyaml, args.outputdir,
    args.isgradsubdir, args.graddir, args.template)
